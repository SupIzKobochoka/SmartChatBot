from docx import Document
from datetime import date 
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .constants import CONSTANTS
from pydantic import BaseModel, Field
from typing import Literal
from langchain.tools import tool
import os

class VacationTransferArgs(BaseModel):
    '''Генерирует файл с заявлением о переносе отпуска'''
    full_name_from: str = Field(description='От кого заявление')
    position_from: str = Field(description='Должность заявителя')
    vacation_from_start: date = Field(description='Дата начала отпуска, который переносится')
    vacation_from_end: date = Field(description='Дата конца отпуска, который переносится')
    vacation_to_start: date = Field(description='Дата начала нового отпуска')
    vacation_to_end: date = Field(description='Дата конца нового отпуска')
    reason_text: str = Field(description='Причина переноса отпуска')

@tool(args_schema=VacationTransferArgs)
def create_vacation_transfer_doc(**data):
    '''Создает файл о переносе отпуска с одной даты на вторую'''
    generate_vacation_transfer_statement(**data)
    return 'Файл о переносе отпуска успешно создан'

def generate_vacation_transfer_statement(
    position_from,
    full_name_from,
    vacation_from_start: date,
    vacation_from_end: date,
    vacation_to_start: date,
    vacation_to_end: date,
    reason_text
):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    today_str = date.today().strftime("%d.%m.%Y")
    title = "ЗАЯВЛЕНИЕ"

    def fmt(d):
        return d.strftime("%d.%m.%Y")

    def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, indent=False):
        p = doc.add_paragraph()
        p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Mm(10)
        for line in text.split("\n"):
            r = p.add_run(line + "\n")
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = bold
        return p

    # Шапка (справа)
    add_paragraph(
        f"В {CONSTANTS['organization']}\n"
        f"{CONSTANTS['position_to']}\n"
        f"{CONSTANTS['full_name_to']}\n\n"
        f"от {position_from}\n"
        f"{full_name_from}",
        align=WD_ALIGN_PARAGRAPH.RIGHT
    )

    doc.add_paragraph()

    # Заголовок
    add_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    doc.add_paragraph()

    # Текст
    body_text = (
        "Прошу перенести ежегодный оплачиваемый отпуск "
        f"с {fmt(vacation_from_start)} по {fmt(vacation_from_end)} "
        f"на период с {fmt(vacation_to_start)} по {fmt(vacation_to_end)}.\n\n"
        f"Причина: {reason_text}."
    )
    add_paragraph(body_text, indent=True)

    doc.add_paragraph("\n")

    # Подпись и дата
    p = doc.add_paragraph()
    p.add_run("Подпись ____________________").font.name = "Times New Roman"
    p.add_run(f"\n\n{today_str}").font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    os.makedirs(CONSTANTS['output_path'](), exist_ok=True)
    doc.save(CONSTANTS['output_path']() / "vacation_transfer_statement.docx")


if __name__ == '__main__':

    generate_vacation_transfer_statement(
        full_name_from="Петрова Петра Петровича",
        position_from="инженера-программиста",
        vacation_from_start=date(2026, 2, 1),
        vacation_from_end=date(2026, 2, 14),
        vacation_to_start=date(2026, 3, 1),
        vacation_to_end=date(2026, 3, 14),
        reason_text="по семейным обстоятельствам"
    )
