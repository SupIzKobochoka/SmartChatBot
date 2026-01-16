from docx import Document
from datetime import date
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .constants import CONSTANTS
from pydantic import BaseModel, Field
from typing import Literal
from langchain.tools import tool
import os

class VacationArgs(BaseModel):
    full_name_from: str = Field(description='От кого заявление')
    position_from: str = Field(description='Должность заявителя')
    body_text: str = Field(description='Основной текст заявления')

@tool(args_schema=VacationArgs)
def create_vacation_doc(**data):
    '''Создает файл об отпуске сотрудника'''
    generate_vacation_statement(**data)
    return 'Файл об отпуске успешно создан'

def generate_vacation_statement(
    position_from,
    full_name_from,
    body_text
):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    date_ = date.today().strftime("%d.%m.%Y")
    title="ЗАЯВЛЕНИЕ"

    def add_paragraph(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, indent=False):
        p = doc.add_paragraph()
        p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Mm(10)
        for line in text.split("\n"):
            r = p.add_run(line + "\n")
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = bold
        return p

    # Шапка (справа)
    add_paragraph(
        f"В {CONSTANTS['organization']}\n"
        f"{CONSTANTS['position_to']}\n"
        f"{CONSTANTS['full_name_to']}\n\n"
        f"от {position_from}\n"
        f"{full_name_from}",
        align=WD_ALIGN_PARAGRAPH.RIGHT
    )

    doc.add_paragraph()

    # Заголовок
    add_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    doc.add_paragraph()

    # Основной текст
    add_paragraph(body_text, indent=True)

    doc.add_paragraph("\n")

    # Подпись и дата
    p = doc.add_paragraph()
    p.add_run("Подпись ____________________").font.name = "Times New Roman"
    p.add_run(f"\n\n{date_}").font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    os.makedirs(CONSTANTS['output_path'](), exist_ok=True)
    doc.save(CONSTANTS['output_path']() / "vacation.docx")


if __name__ == '__main__':
    generate_vacation_statement(
        full_name_from="Петрова Петра Петровича",
        position_from="инженера-программиста",
        body_text=(
            "Прошу предоставить мне ежегодный оплачиваемый отпуск "
            "с 1 февраля 2026 года сроком на 14 календарных дней.\n\n"
            "Причина: по семейным обстоятельствам."
        )
    )